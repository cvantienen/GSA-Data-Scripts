import os
import polars as pl
from docx import Document
from test.sampleCompany import get_sample_company


class SamplePriceComp:
    def __init__(self, conn, contract_number):
        self.conn = conn
        self.company =  get_sample_company(contract_number)
        self.df = None
        self.manufacture_analysis_df = None
        self.analysis_results = None
        self.output_path = "/app/output/"     # Output directory


    def get_100_random_products(self):
        """ Get 100 random items from the specific contract from the database and load it into a DataFrame."""
        query = f"""
        WITH reference_items AS (
            SELECT *
            FROM gsa_product_extract_jan2024
            WHERE contract_number =  '{self.company.contract_number}'
            ORDER BY RANDOM()  -- Randomize the selection of the 100 items
            LIMIT 100
        ),
        competitor_items AS (
            -- Get all items with the same manufacturer_part_number from different contracts
            SELECT *
            FROM gsa_product_extract_jan2024 gi
            WHERE gi.contract_number != '{self.company.contract_number}'  -- Ensure items are from different contracts
            AND gi.manufacturer_part_number IN (SELECT manufacturer_part_number FROM reference_items)
        )
        -- Combine reference items with their matches
        SELECT ref.*, 'reference' AS source
        FROM reference_items ref
        UNION ALL
        SELECT comp.*, 'competitor' AS source
        FROM competitor_items comp
        ORDER BY jprod_id;
        """
        self.df = pl.read_database(query, self.conn)
        self.df = self.df.with_columns(pl.col("price").cast(pl.Float64))

    def analyze_data(self):
        """Analyze the DataFrame and store the results."""
        df = self.df

        # Calculate the comp average price for each manufacturer_part_number
        # for competitor products only
        comp_average_price = df.filter(pl.col("source") == "competitor").group_by("manufacturer_part_number").agg(
            pl.col("price").mean().alias("average_price_on_gsa")
        )

        # Calculate the standard deviation of prices for each manufacturer_part_number
        price_deviation = df.group_by("manufacturer_part_number").agg(
            pl.col("price").std().alias("price_deviation")
        )

        # Combine the average price and price deviation into a single DataFrame
        price_comparison = comp_average_price.join(price_deviation, on="manufacturer_part_number", how="left")

        # Select the columns to include in the final report from the original contractors items
        selected_columns = df.select([
            "contractor_name",
            "contract_number",
            "manufacturer_part_number",
            "manufacturer_name",
            "product_name",
            "price",
        ]).filter(pl.col("contract_number") == self.company.contract_number)

        # Combine the selected columns with the calculated columns
        combined_df = selected_columns.join(price_comparison, on="manufacturer_part_number", how="left")

        # Calculate the percent difference from the Contractors price vs the average price on GSA
        combined_df = combined_df.with_columns(
            ((pl.col("price") - pl.col("average_price_on_gsa")) / pl.col("average_price_on_gsa")).alias("percent_difference")
        )

        # Calculate general statements based on the Price % difference From GSA Average
        below_competitor = combined_df.filter(pl.col("percent_difference") < 0).shape[0]
        above_competitor = combined_df.filter(pl.col("percent_difference") > 0).shape[0]
        avg_percent_diff = combined_df.select(pl.col("percent_difference").mean()).item()
        max_percent_diff = combined_df.select(pl.col("percent_difference").max()).item()
        min_percent_diff = combined_df.select(pl.col("percent_difference").min()).item()

        # Calculate Deviation
        avg_price_deviation = combined_df.select(pl.col("price_deviation").mean()).item()
        max_price_deviation = combined_df.select(pl.col("price_deviation").max()).item()
        min_price_deviation = combined_df.select(pl.col("price_deviation").min()).item()

        # Count items that price deviation is more than $1, $10, and $100
        count_more_than_1 = combined_df.filter(pl.col("price_deviation") > 1).shape[0]
        count_more_than_10 = combined_df.filter(pl.col("price_deviation") > 10).shape[0]
        count_more_than_100 = combined_df.filter(pl.col("price_deviation") > 100).shape[0]

        # Group by manufacturer_name and calculate the average percent difference
        manufacturer_percent_diff = combined_df.group_by("manufacturer_name").agg(
            pl.col("percent_difference").mean().alias("average_percent_difference")
        )

        self.manufacture_analysis_df = manufacturer_percent_diff

        self.analysis_results = {
            "below_competitor": below_competitor,
            "above_competitor": above_competitor,
            "avg_percent_diff": avg_percent_diff,
            "max_percent_diff": max_percent_diff,
            "min_percent_diff": min_percent_diff,
            "avg_price_deviation": avg_price_deviation,
            "max_price_deviation": max_price_deviation,
            "min_price_deviation": min_price_deviation,
            "count_more_than_1": count_more_than_1,
            "count_more_than_10": count_more_than_10,
            "count_more_than_100": count_more_than_100,
        }

    def generate_report(self):
        """Generate a Word document report based on the analysis results."""
        # Create a Word document
        doc = Document()
        doc.add_heading(f'Sample GSA Contract Analysis', 0)
        doc.add_heading(f'GSA Contractor: {self.company}')
        doc.add_heading(f'Contract Number: {self.company.contract_number}', level=1)
        doc.add_heading(f'Manufacturer: {self.company.url}', level=1)
        doc.add_heading(f'Manufacturer: {self.company.url}', level=1)


        results = self.analysis_results
        doc.add_paragraph(f"- Number of items below competitor price: {results['below_competitor']}")
        doc.add_paragraph(f"- Number of items above competitor price: {results['above_competitor']}")
        doc.add_paragraph(f"- Average percent difference: {results['avg_percent_diff']:.2f}%")
        doc.add_paragraph(f"- Max percent difference: {results['max_percent_diff']:.2f}%")
        doc.add_paragraph(f"- Min percent difference: {results['min_percent_diff']:.2f}%")
        doc.add_paragraph(f"- Average price deviation: ${results['avg_price_deviation']:.2f}")
        doc.add_paragraph(f"- Max price deviation: ${results['max_price_deviation']:.2f}")
        doc.add_paragraph(f"- Min price deviation: ${results['min_price_deviation']:.2f}")
        doc.add_paragraph(f"- Count of items with price deviation more than $1: {results['count_more_than_1']}")
        doc.add_paragraph(f"- Count of items with price deviation more than $10: {results['count_more_than_10']}")
        doc.add_paragraph(f"- Count of items with price deviation more than $100: {results['count_more_than_100']}")

        # Add a section for the manufacturer-specific analysis
        # Filter out empty columns
        non_empty_columns = [col for col in self.manufacture_analysis_df.columns if
                             not self.manufacture_analysis_df[col].is_null().all()]

        filtered_df = self.manufacture_analysis_df.select(non_empty_columns)

        table = doc.add_table(rows=1, cols=len(filtered_df.columns))
        header_cells = table.rows[0].cells
        for i, column in enumerate(filtered_df.columns):
            header_cells[i].text = column

        for row in filtered_df.to_dicts():
            row_cells = table.add_row().cells
            for i, (key, value) in enumerate(row.items()):
                row_cells[i].text = str(value)


        # Save the document
        doc.save(f"SampleReport_{self.company.contract_number}.docx")

        # Construct the full file path
        file_name = f"SampleReport_{self.company.contract_number}.docx"
        full_path = os.path.join(self.output_path, file_name)

        # Save the document
        doc.save(full_path)

